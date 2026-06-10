# -*- coding: utf-8 -*-
"""
生成「2026 小X宝开源医疗社区黑客松」赛事详情页最终文案排版 DOCX
包含 Banner 图文排版、赛事信息、评审规则、奖励体系、提交规范、FAQ 等
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Paths ──────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR = os.path.dirname(SCRIPT_DIR)
BANNER1 = os.path.join(SCRIPT_DIR, "Banner样式1.png")
BANNER2 = os.path.join(SCRIPT_DIR, "Banner样式2.png")
POSTER = os.path.join(SCRIPT_DIR, "宣传海报样本.png")
LOGO_XIAOXBAO = os.path.join(ROOT_DIR, "小X宝社区logo元素", "小X宝.png")
LOGO_MODA = os.path.join(ROOT_DIR, "赞助商logo", "魔搭logo.png")
LOGO_KNOWS = os.path.join(ROOT_DIR, "赞助商logo", "KnowS logo.png")
LOGO_STEPFUN = os.path.join(ROOT_DIR, "赞助商logo", "阶跃logo.png")

OUTPUT = os.path.join(SCRIPT_DIR, "赛事详情页文案排版_终稿.docx")

# ── Theme Colors ──────────────────────────────────
ORANGE_BRAND = RGBColor(0xFF, 0x6B, 0x35)    # 品牌橙
GREEN_BRAND  = RGBColor(0x3D, 0xB8, 0x8C)    # 品牌绿
DARK_TEXT     = RGBColor(0x2D, 0x2D, 0x2D)    # 深灰正文
GRAY_TEXT     = RGBColor(0x66, 0x66, 0x66)    # 辅助灰
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helpers ────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """给表格单元格设置底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_paragraph(doc, text, font_name="微软雅黑", font_size=11,
                          bold=False, color=None, alignment=None, space_after=6,
                          space_before=0):
    """添加带完整样式的段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def add_section_heading(doc, text, level=1):
    """添加带品牌色的章节标题"""
    size_map = {1: 18, 2: 14, 3: 12}
    p = add_styled_paragraph(
        doc, text,
        font_size=size_map.get(level, 14),
        bold=True,
        color=ORANGE_BRAND,
        space_before=18,
        space_after=10
    )
    # 加下划线装饰
    if level == 1:
        doc.add_paragraph("━" * 40).runs[0].font.color.rgb = RGBColor(0xE0, 0xE0, 0xE0)
    return p

def add_bullet(doc, text, bold_prefix=None, indent_level=0):
    """添加项目符号段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Cm(1.0 + indent_level * 0.8)

    # 符号
    marker = "◆ " if indent_level == 0 else "• "
    run_marker = p.add_run(marker)
    run_marker.font.name = "微软雅黑"
    run_marker.font.size = Pt(10)
    run_marker.font.color.rgb = GREEN_BRAND

    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.name = "微软雅黑"
        run_b._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
        run_b.font.size = Pt(10.5)
        run_b.font.bold = True
        run_b.font.color.rgb = DARK_TEXT

    run_t = p.add_run(text)
    run_t.font.name = "微软雅黑"
    run_t._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    run_t.font.size = Pt(10.5)
    run_t.font.color.rgb = DARK_TEXT
    return p

def add_info_table(doc, rows_data, col_widths=None):
    """添加信息表格"""
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
            run.font.size = Pt(10)

            if i == 0:  # 表头
                run.font.bold = True
                run.font.color.rgb = WHITE
                set_cell_shading(cell, "3DB88C")
            else:
                run.font.color.rgb = DARK_TEXT
                if i % 2 == 0:
                    set_cell_shading(cell, "F5FAF7")

            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 设置表格边框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    return table


# ══════════════════════════════════════════════════
#  Main Document
# ══════════════════════════════════════════════════

doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# ── 设置默认字体 ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
font.color.rgb = DARK_TEXT
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ═══════════════════════════════════════════
# Banner 区域
# ═══════════════════════════════════════════

if os.path.exists(BANNER2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(BANNER2, width=Cm(17))
    p.paragraph_format.space_after = Pt(6)

# ── 标题区 ──
add_styled_paragraph(
    doc,
    "2026 小X宝开源医疗社区黑客松",
    font_size=24, bold=True, color=ORANGE_BRAND,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=12, space_after=4
)

add_styled_paragraph(
    doc,
    "光已成炬，照亮崎岖 | Light Turns Into Torches, Illuminating the Rugged Path",
    font_size=12, bold=False, color=GREEN_BRAND,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=0, space_after=2
)

add_styled_paragraph(
    doc,
    "小X宝开源医疗社区 × 魔搭 ModelScope 联合主办",
    font_size=11, bold=False, color=GRAY_TEXT,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=0, space_after=16
)

# ═══════════════════════════════════════════
# 一、竞赛简介
# ═══════════════════════════════════════════

add_section_heading(doc, "一、竞赛简介")

add_styled_paragraph(
    doc,
    "在医疗的广阔疆域中，仍有许多崎岖之路——肿瘤的复杂、罕见病的孤独。"
    "科技的力量，应当成为照亮这些角落的火炬。",
    font_size=11, color=DARK_TEXT, space_after=8
)

add_styled_paragraph(
    doc,
    "2026 小X宝开源医疗社区黑客松 联合 ModelScope 魔搭社区正式拉开帷幕！"
    "我们邀请每一位心怀善意的开发者，用 AI 技能（Skills）与 MCP 工具，为生命构筑更多可能。",
    font_size=11, color=DARK_TEXT, space_after=8
)

add_styled_paragraph(
    doc,
    "本次黑客松由小X宝开源医疗公益社区发起，旨在通过开源技术服务真实医疗需求。"
    "无论你是算法极客、医疗 AI 实践者，还是对 Agent 领域充满好奇的新手，这里都有你的舞台。",
    font_size=11, color=DARK_TEXT, space_after=12
)

# 合作方信息
add_info_table(doc, [
    ["角色", "名称", "说明"],
    ["联合主办", "魔搭 ModelScope", "阿里旗下 AI 开发者社区，提供赛事平台与运营资源"],
    ["合作方", "KnowS", "参与联合推广"],
    ["大模型赞助", "阶跃星辰 StepFun", "提供大模型 API 额度支持"],
])

doc.add_paragraph()  # spacing

# ═══════════════════════════════════════════
# 二、赛题方向与技术方案
# ═══════════════════════════════════════════

add_section_heading(doc, "二、赛题方向与技术方案")

add_styled_paragraph(
    doc,
    "赛题方向为通用医学 + 生命科学方向，不做更细的限制。"
    "核心赛题：聚焦医疗垂直领域，构建可复用的 Skills 或 MCP 扩展工具。\n"
    "赛事特色：不设细分赛道，只要你的应用能解决真实医疗场景问题，即刻出发！",
    font_size=11, color=DARK_TEXT, space_after=8
)

add_styled_paragraph(doc, "作品须满足以下条件：", font_size=11, bold=True, color=DARK_TEXT, space_after=4)

add_bullet(doc, "聚焦医疗垂直领域", bold_prefix="医疗场景导向：")
add_bullet(doc, "（如肿瘤、罕见病、诊断辅助、患者管理、医学文献检索等）。", indent_level=1)
add_bullet(doc, "形式为 MCP 工具 或 Agent Skill，可独立运行或集成到现有 Agent 框架。", bold_prefix="技术形式：")
add_bullet(doc, "须能在魔搭社区部署为 Model/Space。", bold_prefix="平台部署：")
add_bullet(doc, "选题不限细分方向，但需要在开发前在群内或活动页面登记选题，避免重复。", bold_prefix="选题登记：")

doc.add_paragraph()

# ═══════════════════════════════════════════
# 三、赛程安排
# ═══════════════════════════════════════════

add_section_heading(doc, "三、赛程安排：紧凑冲刺，直通 WAIC")

add_styled_paragraph(
    doc,
    "本次大赛为期近四周（6月18日 - 7月12日），节奏紧凑，直达 WAIC 世界人工智能大会：",
    font_size=11, color=DARK_TEXT, space_after=8
)

add_info_table(doc, [
    ["阶段", "时间", "内容"],
    ["上线 / 报名开始", "2026年6月18日", "在魔搭社区完成身份绑定，加入官方交流群，即刻开发"],
    ["选题与开发", "2026年6月18日 - 7月11日", "确定医疗选题并开始开发，期间可随时在魔搭提交版本"],
    ["作品提交截止", "2026年7月12日 23:59", "将作品发布至魔搭社区，并填写最终提交表单"],
    ["集中评审", "2026年7月13日 - 7月14日", "周末双轨制评分（社区热度 40% + 专家评审 60%）"],
    ["结果公示与颁奖", "2026年7月15日", "公布获奖名单，Top3 获奖者发放 WAIC 门票"],
    ["WAIC 赛果展示", "2026年7月17日 - 7月20日", "Top3 获奖者于上海 WAIC 大会现场做赛果展示"],
])

doc.add_paragraph()

# ═══════════════════════════════════════════
# 四、评审规则
# ═══════════════════════════════════════════

add_section_heading(doc, "四、评审规则（双轨加权评分制）")

add_styled_paragraph(
    doc,
    "我们采用「社区热度 + 专家评审」的双轨评分制，确保每一份优秀作品都能被看见：",
    font_size=11, color=DARK_TEXT, space_after=8
)

add_info_table(doc, [
    ["评分维度", "权重", "评分依据"],
    ["社区热度分", "40%", "基于作品在魔搭社区的调用数与收藏数"],
    ["专家评审分", "60%", "由医疗专家与技术大咖从实用性、代码质量、文档完整性三大维度打分"],
])

doc.add_paragraph()

# ═══════════════════════════════════════════
# 五、奖励体系
# ═══════════════════════════════════════════

add_section_heading(doc, "五、奖励体系")

add_styled_paragraph(
    doc,
    "（注：以下奖励方案为初拟版本，最终金额与形式以与魔搭社区协商结果为准）",
    font_size=9, color=GRAY_TEXT, space_after=8
)

add_info_table(doc, [
    ["奖项", "名额", "奖励内容"],
    ["🏆 一等奖", "Top 3（3名）", "1000 RMB 云资源或 API 额度 + WAIC 门票"],
    ["🥈 二等奖", "Top 4-10（7名）", "500 RMB 云资源或 API 额度"],
    ["🥉 三等奖", "Top 11-20（10名）", "100 RMB 云资源或 API 额度"],
])

doc.add_paragraph()

add_styled_paragraph(doc, "额外权益：", font_size=11, bold=True, color=DARK_TEXT, space_after=4)
add_bullet(doc, "获胜者将赢取丰厚的云资源或 API 额度，助力项目持续演进。", bold_prefix="资源支持：")
add_bullet(doc, "所有完成提交的参与者均可获得官方电子参与证书。", bold_prefix="荣誉见证：")
add_bullet(doc, "优秀作品将上架 GitHub 与魔搭社区，获得全平台流量支持。", bold_prefix="品牌曝光：")

doc.add_paragraph()

# ═══════════════════════════════════════════
# 六、作品提交规范
# ═══════════════════════════════════════════

add_section_heading(doc, "六、作品提交规范")

add_styled_paragraph(
    doc,
    "参赛者在魔搭活动页填写表单时，必须包含以下三个核心交付物：",
    font_size=11, color=DARK_TEXT, space_after=8
)

add_bullet(doc, "（部署在魔搭的 Skill 或 MCP，确保公开可访问，有基本 Demo）", bold_prefix="魔搭社区作品链接 ")
add_bullet(doc, "（如 GitHub、Gitee 等，需含开源 License，推荐 Apache 2.0 或 MIT）", bold_prefix="公开代码仓库地址 ")
add_bullet(doc, "（须包含：背景说明、医疗场景解决点、使用方法、部署步骤）", bold_prefix="规范 README 文档 ")

doc.add_paragraph()

add_styled_paragraph(doc, "提交入口：", font_size=11, bold=True, color=DARK_TEXT, space_after=4)
add_bullet(doc, "https://modelscope.cn/skills/create?template=custom", bold_prefix="Skill 提交：")
add_bullet(doc, "https://modelscope.cn/mcp/servers/create", bold_prefix="MCP 提交：")

doc.add_paragraph()

# ── 仓库设置规范 ──
add_section_heading(doc, "七、仓库设置规范", level=2)

add_bullet(doc, "仓库必须设置为 Public（公开）。", bold_prefix="可见性：")
add_bullet(doc, "仓库根目录必须包含 LICENSE 文件。建议使用 MIT 或 Apache 2.0 协议，以鼓励社区二次开发。", bold_prefix="开源协议：")
add_bullet(doc, "敏感信息（如 API Keys、个人隐私数据）绝不可硬编码提交，需使用 .env 等方式隔离。", bold_prefix="代码整洁：")
add_bullet(doc, "禁止使用真实患者数据，不得做出诊断承诺。", bold_prefix="医疗合规：")

doc.add_paragraph()

# ── README 模板 ──
add_section_heading(doc, "八、README.md 模板要求", level=2)

add_styled_paragraph(
    doc,
    "README 是评委了解作品的第一窗口，必须包含以下结构：",
    font_size=11, color=DARK_TEXT, space_after=6
)

readme_template = """# [项目名称]

## 1. 项目简介与医疗场景
- **一句话描述**：[用一句话说明这是一个什么医疗 Skill/MCP]
- **解决的痛点**：[例如：肿瘤患者随访数据难结构化、罕见病文献检索效率低等]
- **目标受众**：[医生、患者、医学研究员等]

## 2. 功能特性
- [特性 1：如支持基于 PubMed 的自动文献检索]
- [特性 2：如支持与外部知识图谱联动]

## 3. 魔搭社区运行/部署指南
- **魔搭展示链接**：[提供您的 Skill/MCP 在魔搭的链接]
- **本地运行步骤**：
  ```bash
  pip install -r requirements.txt
  python main.py
  ```

## 4. 演示与输入输出示例
- **输入示例**：`[用户的 Prompt 或 API 请求格式]`
- **预期输出**：`[预期的结果格式或截图]`
  *(强烈建议附上 1-2 张运行结果截图，或演示视频链接)*

## 5. 局限性与未来规划
- 目前版本存在的不足（请务必如实说明，医疗应用需严谨）
- 未来拟加入的新功能

## 6. 团队与致谢
- 成员介绍及分工
- 致谢使用到的开源项目或数据集"""

p = doc.add_paragraph()
run = p.add_run(readme_template)
run.font.name = 'Consolas'
run.font.size = Pt(8.5)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
p.paragraph_format.space_after = Pt(8)

doc.add_paragraph()

# ── 加分项 ──
add_section_heading(doc, "九、专家评委额外加分项（Bonus）", level=2)

add_bullet(doc, "明确说明该工具的输出不作为最终医疗诊断依据。", bold_prefix="医疗合规性声明：")
add_bullet(doc, "提供该 Skill/MCP 在特定测试集上的准确率或可用性评估。", bold_prefix="评测数据报告：")
add_bullet(doc, "提供一段 1-3 分钟的视频，展示完整工作流。", bold_prefix="演示视频：")

doc.add_paragraph()

# ═══════════════════════════════════════════
# 十、常见问题 FAQ
# ═══════════════════════════════════════════

add_section_heading(doc, "十、常见问题 FAQ")

faq_list = [
    ("Q1：赛事详情页采用什么形式？",
     "采用图文详情页，方便实时调整文案。主视觉及 Banner 已完成定稿。"),
    ("Q2：谁可以参加？",
     "面向全球开发者、医疗 AI 实践者与开源社区成员，不限学历、不限职业，个人或团队均可参加。"),
    ("Q3：可以用什么技术栈？",
     "不限技术栈，但作品形式须为 MCP 工具或 Agent Skill，且须能在魔搭社区部署为 Model/Space。"),
    ("Q4：提交后会审核吗？",
     "提交后只要不是反动信息被安全拦截，都会发布在社区里。赛后统一整理为赛果合集进行宣传。"),
    ("Q5：选题可以重复吗？",
     "建议在开发前于群内或活动页面登记选题，避免重复。不同角度的实现不受限制。"),
    ("Q6：奖励如何发放？",
     "云资源或 API 额度将在7月15日结果公示后统一发放，Top3 获奖者同时获得 WAIC 世界人工智能大会（7月17-20日，上海）入场门票。"),
    ("Q7：如何加入交流群？",
     "报名成功后将获得官方微信群入口，也可扫描宣传海报上的二维码加入。"),
]

for q, a in faq_list:
    add_styled_paragraph(doc, q, font_size=11, bold=True, color=ORANGE_BRAND, space_before=8, space_after=2)
    add_styled_paragraph(doc, a, font_size=10.5, color=DARK_TEXT, space_after=6)

# ═══════════════════════════════════════════
# 宣传海报
# ═══════════════════════════════════════════

add_section_heading(doc, "附：宣传海报")

if os.path.exists(POSTER):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(POSTER, width=Cm(12))
    p.paragraph_format.space_after = Pt(12)

# ═══════════════════════════════════════════
# 页脚信息
# ═══════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph("━" * 40).runs[0].font.color.rgb = RGBColor(0xE0, 0xE0, 0xE0)

add_styled_paragraph(
    doc,
    "主办方：小X宝开源医疗社区 × 魔搭 ModelScope",
    font_size=10, color=GRAY_TEXT,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=2
)
add_styled_paragraph(
    doc,
    "合作方：KnowS | 大模型赞助：阶跃星辰 StepFun",
    font_size=10, color=GRAY_TEXT,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=2
)
add_styled_paragraph(
    doc,
    "项目默认开源 · 禁真实患者数据 · 不做诊断承诺",
    font_size=9, color=GRAY_TEXT,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=2
)
add_styled_paragraph(
    doc,
    f"文档生成日期：2026年6月10日（时间节点更新版）",
    font_size=9, color=RGBColor(0xAA, 0xAA, 0xAA),
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=0
)

# ── Save ──
doc.save(OUTPUT)
print(f"[OK] Document saved: {OUTPUT}")
